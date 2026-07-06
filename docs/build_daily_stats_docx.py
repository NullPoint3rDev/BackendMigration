"""Генерирует Word-документ с алгоритмами суточной статистики.

Запуск: python3 build_daily_stats_docx.py
Результат: daily-stats-algorithms.docx рядом со скриптом.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CODE_FONT = "Consolas"
BODY_FONT = "Calibri"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
CODE_BG = "F2F2F2"


def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def code_block(doc, lines):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        shade(p, CODE_BG)
        run = p.add_run(line)
        run.font.name = CODE_FONT
        run.font.size = Pt(10)


def add_body(doc, runs):
    """runs: список кортежей (текст, bold, mono)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for text, bold, mono in runs:
        r = p.add_run(text)
        r.bold = bold
        if mono:
            r.font.name = CODE_FONT
            r.font.size = Pt(10)
    return p


def bullet(doc, runs, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    for text, bold, mono in runs:
        r = p.add_run(text)
        r.bold = bold
        if mono:
            r.font.name = CODE_FONT
            r.font.size = Pt(10)
    return p


def numbered(doc, runs):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    for text, bold, mono in runs:
        r = p.add_run(text)
        r.bold = bold
        if mono:
            r.font.name = CODE_FONT
            r.font.size = Pt(10)
    return p


def T(text):
    return (text, False, False)


def B(text):
    return (text, True, False)


def M(text):
    return (text, False, True)


doc = Document()

# Базовый стиль
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)

# Заголовок
title = doc.add_heading("Алгоритмы суточной статистики сварочного аппарата", level=0)

intro = doc.add_paragraph()
intro.add_run(
    "Расчёты реализованы в WeldingMachineDailyStatsService (пересчёт суток "
    "recomputeDayAndSave) с использованием WeldingStateDurationUtil и "
    "MonitorActivityClassifier."
).italic = True

# --- 1 ---
doc.add_heading("1. Расчёт таймеров активности за сутки", level=1)
add_body(doc, [
    T("Считаются четыре таймера (в миллисекундах): "),
    M("off"), T(" (выключен), "),
    M("error"), T(" (авария/ошибка), "),
    M("on"), T(" (включён/дежурит), "),
    M("welding"), T(" (сварка)."),
])

add_body(doc, [
    B("Границы суток. "),
    T("Начало суток — "), M("statDate"), T(" в 00:01 ("), M("dayStart"),
    T(", не с полуночи). Конец — начало следующих суток по таймзоне "),
    M("Europe/Moscow"), T(" ("), M("dayEnd"),
    T("). Если считаются текущие сутки, конец обрезается до текущего момента ("),
    M("effectiveEnd = min(now, dayEnd)"), T(")."),
])

add_body(doc, [
    B("Загрузка данных. "),
    T("Из БД берутся все состояния аппарата ("), M("WeldingMachineState"),
    T(") за интервал "), M("[dayStart, dayEnd)"),
    T(", сортируются по "), M("dateCreated"),
    T(". Для каждого состояния подгружаются его параметры: текст состояния "
      "(«Состояние аппарата» / "), M("WeldingMachineState"), T(" / "),
    M("State.WeldingMachineState"), T(") и ток ("), M("State.I"),
    T(" / «Ток» / "), M("Current"), T(")."),
])

add_body(doc, [B("Проход по состояниям. "), T("Для каждого состояния:")])

numbered(doc, [
    T("Вычисляется пересечение его длительности с окном "),
    M("[dayStart, effectiveEnd)"), T(" — функция "), M("overlapDurationMs"), T(":"),
])
bullet(doc, [T("длительность состояния = "), M("stateDurationMs"),
             T(", если задана; иначе = зазор до "), M("dateCreated"),
             T(" следующего состояния;")], level=1)
bullet(doc, [T("у последнего состояния хвост до "), M("now"),
             T(" намеренно не добавляется ("), M("openEndIfLast = null"),
             T(") — иначе таймеры «скачут» на последнем опросе;")], level=1)
bullet(doc, [T("берётся длина именно перекрытия отрезка состояния с окном суток;")], level=1)
bullet(doc, [T("если перекрытие ≤ 0 — состояние пропускается.")], level=1)

numbered(doc, [
    T("Состояние классифицируется в один из режимов ("),
    M("MonitorActivityClassifier.classify"), T(") по приоритету:"),
])
bullet(doc, [B("welding"), T(" — текст содержит «сварка»/welding/weld/«сварочн», "
             "ИЛИ статус = Welding, ИЛИ текста нет, но ток > 1 А;")], level=1)
bullet(doc, [B("error"), T(" — текст содержит «авария»/«ошибка»/error/emergency/failure, "
             "ИЛИ errorCode > 0 (или нечисловой), ИЛИ статус = Error;")], level=1)
bullet(doc, [B("off"), T(" — текст содержит «выключ»/off/offline/«не в сети»/«дежур»/standby, "
             "ИЛИ статус = Offline;")], level=1)
bullet(doc, [B("on"), T(" — «ожидан»/waiting/«включ»/idle/ready, и как режим по умолчанию.")], level=1)

numbered(doc, [T("Длительность перекрытия прибавляется к соответствующему таймеру.")])

add_body(doc, [
    B("Замощение суток. "),
    T("Чтобы сумма таймеров равнялась прошедшему времени "),
    M("[00:01, effectiveEnd]"),
    T(", непокрытое телеметрией время относится к "), M("off"),
    T(" («Выкл./не в сети»):"),
])
bullet(doc, [T("голова "), M("[00:01 → первый пакет дня]"), T(";")], level=1)
bullet(doc, [T("хвост "), M("[последний покрытый момент → effectiveEnd]"),
             T(" (последний покрытый момент = "), M("dateCreated"),
             T(" последнего состояния + его "), M("stateDurationMs"),
             T("; хвост до "), M("now"),
             T(" намеренно не тянется, иначе таймеры «скачут»);")], level=1)
bullet(doc, [T("весь интервал суток, если пакетов не было вовсе.")], level=1)
add_body(doc, [
    T("Штатные короткие зазоры между опросами уже учтены через «зазор до следующего "
      "состояния»."),
])

add_body(doc, [
    B("Итог. "), T("Суммы "), M("offMs / errorMs / onMs / weldingMs"),
    T(" сохраняются в "), M("WeldingMachineDailyStats"),
    T(". В DTO поле "), M("standbyMs"), T(" из БД отдаётся как "), M("errorMs"),
    T(" (переименование), а "), M("standbyMs"), T(" наружу = 0."),
])

# --- 2 ---
doc.add_heading("2. Расчёт суточного расхода газа", level=1)
add_body(doc, [
    T("Расход газа считается по накопительному счётчику "),
    M("Core.GasConsumptionSincePowerOn"),
    T(" (литры «с момента включения») методом суммирования положительных дельт."),
])

add_body(doc, [
    B("Данные. "),
    T("Берутся состояния за расширенное окно "), M("[dayStart − 1 день, dayEnd)"),
    T(" (лишние сутки нужны, чтобы найти «базовое» значение счётчика на начало суток). "
      "Для каждого состояния подгружается значение накопительного счётчика газа."),
])

add_body(doc, [B("Проход по состояниям (по возрастанию времени):")])
bullet(doc, [T("если состояние раньше "), M("dayStart"),
             T(" — его значение запоминается как "), M("baselineAtDayStart"),
             T(" и как "), M("lastCumulative"),
             T(" (накоплено до начала суток), в сумму не идёт;")])
bullet(doc, [T("если состояние в пределах суток:")])
bullet(doc, [T("если baseline ещё не задан — первое значение становится baseline;")], level=1)
bullet(doc, [T("к сумме прибавляется дельта между "), M("lastCumulative"),
             T(" и текущим значением ("), M("sumGasCumulativeDelta"), T("):")], level=1)
bullet(doc, [T("если счётчик вырос ("), M("current ≥ last"),
             T(") → дельта = "), M("current − last"), T(";")], level=1)
bullet(doc, [T("если счётчик упал, но это распознано как сброс питания → дельта = само "),
             M("current"), T(" (счётчик обнулился и начал заново);")], level=1)
bullet(doc, [T("иначе (мелкая просадка от обрыва связи) → дельта = 0;")], level=1)
bullet(doc, [M("lastCumulative"),
             T(" обновляется до "), M("current"),
             T(" только при росте или реальном сбросе (иначе провал игнорируется).")], level=1)

add_body(doc, [
    B("Определение сброса счётчика "), M("(isGasCounterPowerOnReset)"),
    T(": считается сбросом, только если падение ≥ 100 л И новое значение < 50% от "
      "предыдущего. Это отсекает мелкие глюки связи от настоящего обнуления при "
      "перезагрузке аппарата."),
])

add_body(doc, [
    B("Итог. "), M("gasConsumptionL"),
    T(" = сумма дельт (округление до 3 знаков), плюс сохраняется "),
    M("gasBaselineAtDayStartL"), T(". Если данных счётчика нет — расход 0."),
])

# --- 3 ---
doc.add_heading("3. Расчёт суточного расхода проволоки", level=1)
add_body(doc, [
    T("Параметр «Расход проволоки» ("), M("WIRE_PARAM"),
    T(") с аппарата — это накопительная длина поданной проволоки в метрах «с момента "
      "включения» (uint32 из пакета Core, декодированный как IEEE-float), а НЕ мгновенная "
      "скорость. Значение монотонно растёт при сварке и стоит в простое, обнуляется при "
      "перезагрузке аппарата. Поэтому расход считается по дельтам счётчика, как газ."),
])

add_body(doc, [
    B("Данные. "),
    T("Состояния за сутки "), M("[dayStart, dayEnd)"),
    T(" по времени. Для каждого подгружается значение счётчика «Расход проволоки» "
      "(обычный маппинг + fallback на нативный "), M("COALESCE"), T(")."),
])

add_body(doc, [
    B("Константа. "),
    T("Линейная плотность проволоки "), M("wireLinearDensityKgPerMeter"),
    T(" = 0.0089 кг/м (сталь Ø1.2 мм, из конфигурации "),
    M("report.wire.linear-density-kg-per-meter"), T(")."),
])

add_body(doc, [
    B("Проход по состояниям (по возрастанию времени), "), M("sumWireCumulativeMeters"),
    T(":"),
])
numbered(doc, [T("Для каждой пары последовательных значений счётчика берётся дельта.")])
numbered(doc, [T("Если счётчик вырос → прибавляем "), M("current − last"), T(".")])
numbered(doc, [T("Если счётчик крупно упал (более чем вдвое, "),
               M("current < last × 0.5"),
               T(") — это перезагрузка «с включения» → прибавляем само "), M("current"),
               T(".")])
numbered(doc, [T("Мелкая просадка (шум/обрыв связи) → игнор, причём "), M("last"),
               T(" вниз не двигаем (иначе шумовой «отскок» даст ложный прирост).")])
numbered(doc, [T("Сумма приростов = поданная за сутки длина проволоки (м).")])

add_body(doc, [T("Формула по сути:")])
code_block(doc, ["масса (кг) = Σ приростов_счётчика (м) × плотность (кг/м)"])

add_body(doc, [
    B("Итог. "),
    T("Метры × плотность, округление до 5 знаков, сохраняется как "),
    M("wireConsumptionKg"),
    T(". Плотность фиксирована под сталь Ø1.2 мм; апгрейд — таблица плотностей по "
      "диаметру/материалу (в пакете есть "), M("weldingWireDiameter"), T("/"),
    M("weldingMaterial"), T(")."),
])

# --- Общие замечания ---
doc.add_heading("Общие замечания", level=1)
add_body(doc, [
    T("Пересчёт ("), M("recomputeDay"), T(") детерминирован и кэшируется в "),
    M("WeldingMachineDailyStats"),
    T(". Для текущих суток при устаревшем кэше (старше 60 с) идёт синхронный пересчёт, "
      "для прошлых дат — асинхронный с дебаунсом 15 с."),
])

import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "daily-stats-algorithms.docx")
doc.save(out)
print("saved:", out)
